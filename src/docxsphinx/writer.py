# -*- coding: utf-8 -*-
"""
    sphinxcontrib-docxwriter
    ~~~~~~~~~~~~~~~~~~~~~~~~~~

    Custom docutils writer for OpenXML (docx).

    :copyright:
        Copyright 2010 by shimizukawa at gmail dot com (Sphinx-users.jp).
    :license: BSD, see LICENSE for details.
"""
from __future__ import division

import logging
import os
import sys

from docutils import nodes, writers
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
# noinspection PyUnresolvedReferences
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Inches
# noinspection PyProtectedMember
from docx.table import _Cell

logging.basicConfig(
    filename='docx.log',
    filemode='w',
    level=logging.INFO,
    format="%(asctime)-15s  %(message)s"
)
logger = logging.getLogger('docx')


def dprint(_func=None, **kw):
    """Print debug information."""
    # noinspection PyProtectedMember
    f = sys._getframe(1)
    if kw:
        text = ', '.join('%s = %s' % (k, v) for k, v in kw.items())
    else:
        text = dict((k, repr(v)) for k, v in f.f_locals.items() if k != 'self')
        text = str(text)

    if _func is None:
        _func = f.f_code.co_name

    if not (_func.startswith("visit") or _func.startswith("depart")) and 'node' in f.f_locals:
        _func = "?_{}".format(f.f_locals['node'].__class__.__name__)

    # It would be nice to have the non-kw dprints to be debug-level' issues,
    # but that does not seem to work.
    if kw:
        logger.info(' '.join([_func, text]))

    logger.info(' '.join([_func, text]))  # HB TODO Remove
    # print(' '.join([_func, text]))  # HB TODO Remove


# noinspection PyUnusedLocal
def _make_depart_admonition(name):
    # noinspection PyMissingOrEmptyDocstring,PyUnusedLocal
    def depart_admonition(self, node):
        dprint()
        raise nodes.SkipNode
        # from sphinx.locale import admonitionlabels, versionlabels, _

    return depart_admonition


# noinspection PyClassicStyleClass,PyMissingOrEmptyDocstring
class DocxWriter(writers.Writer):
    """docutil writer class for docx files"""
    supported = ('docx',)
    settings_spec = ('No options here.', '', ())
    settings_defaults = {}

    output = None
    template_dir = "NO"

    def __init__(self, builder):
        writers.Writer.__init__(self)
        self.builder = builder
        self.template_setup()  # setup before call almost docx methods.

        if self.template_dir == "NO":
            dc = Document()
        else:
            dc = Document(os.path.join('source', self.template_dir))
        self.docx_container = dc

    def template_setup(self):
        dotx = self.builder.config['docx_template']
        if dotx:
            logger.info("MK using template {}".format(dotx))
            self.template_dir = dotx

    def save(self, filename):
        self.docx_container.save(filename)

    def translate(self):
        visitor = DocxTranslator(
            self.document, self.builder, self.docx_container)
        self.document.walkabout(visitor)
        self.output = ''  # visitor.body


class DocxState(object):
    """
    DocxState class keeps track of which part of the document is being worked on.

    In particular it is used to allow lists in tables.
    """

    def __init__(self, location=None):
        self.location = location
        self.table = None
        self.column_widths = None
        self.table_style = None
        self.more_cols = 0
        self.row = None
        self.cell_counter = 0
        self.next_figure_num = 1
        "Next figure number to assign"
        self.ncolumns = 1
        "Number of columns in the current table."


# noinspection PyClassicStyleClass,PyMissingOrEmptyDocstring,PyUnusedLocal
class DocxTranslator(nodes.NodeVisitor):
    """Visitor class to create docx content."""

    def __init__(self, document, builder, docx_container):
        self.builder = builder
        self.docx_container = docx_container
        nodes.NodeVisitor.__init__(self, document)

        # TODO: Perhaps move the list_style into DocxState.
        # However, it should still be a list, and not a separate state,
        # because nested lists are not really nested.
        # So it will only be necessary if there are lists in tables
        # that are in lists.
        self.list_style = []
        self.list_level = 0

        # TODO: And what about sectionlevel?
        self.sectionlevel = 0

        self.table_style_default = 'Medium Grid 1 Accent 1'
        self.in_literal_block = False
        self.in_figure = False
        self.strong = False
        self.emphasis = False
        self.center = False

        self.current_state = DocxState(location=self.docx_container)
        self.current_state.table_style = self.table_style_default

        "The place where paragraphs will be added."
        self.old_states = []
        "A list of older states, e.g. typically [document, table-cell]"

        self.current_paragraph = self.current_state.location.add_paragraph("")
        "The current paragraph that text is being added to."

    def add_text(self, text):
        dprint()
        textrun = self.current_paragraph.add_run(text)
        if self.strong:
            textrun.bold = True
        if self.emphasis:
            textrun.italic = True

    def add_paragraph(self, dest, text='', style=None):
        p = dest.add_paragraph(text, style)

        if self.center:
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        return p

    def add_seq_field(self, field_type, contents):
        paragraph = self.current_paragraph
        run = paragraph.add_run()
        r = run._r

        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        r.append(fldChar)

        run = paragraph.add_run()
        r = run._r

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        logger.info(' SEQ {} \* ARABIC'.format(field_type))
        instrText.text = ' SEQ {} \* ARABIC '.format(field_type)
        r.append(instrText)

        run = paragraph.add_run()
        r = run._r

        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'separate')
        r.append(fldChar)

        paragraph.add_run(str(contents))

        run = paragraph.add_run()
        r = run._r

        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'end')
        r.append(fldChar)

    def new_state(self, location):
        dprint()
        self.old_states.append(self.current_state)
        self.current_state = DocxState(location=location)

    def end_state(self, first=None):
        dprint()
        self.current_state = self.old_states.pop()

    def print_and_skip(self, node):
        dprint()
        raise nodes.SkipNode

    def just_print(self, node):
        dprint()
        pass

    def visit_start_of_file(self, node):
        dprint()
        # TODO: HB should visit_start_of_file reset the sectionlevel?
        # If so, should it start a new state? If so, with which location?

        # FIXME: visit_start_of_file not close previous section.
        # sectionlevel keep previous and new file's heading level start with
        # previous + 1.
        # This quick hack reset sectionlevel per file.
        # (BTW Sphinx has heading levels per file? or entire document?)
        self.sectionlevel = 0

    def visit_comment(self, node):
        dprint()
        # TODO: FIX Dirty hack / kludge to set table style.
        # Use proper directives or something like that
        comment = node[0]
        if 'DocxTableStyle' in comment:
            self.current_state.table_style = comment.split('DocxTableStyle')[
                -1].strip()
        raise nodes.SkipNode

    def visit_section(self, node):
        dprint()
        self.sectionlevel += 1

    def depart_section(self, node):
        dprint()
        if self.sectionlevel > 0:
            self.sectionlevel -= 1

    def visit_strong(self, node):
        dprint()
        # self.add_text('**')
        self.strong = True

    def depart_strong(self, node):
        dprint()
        # self.add_text('**')
        self.strong = False

    def visit_emphasis(self, node):
        dprint()
        # self.add_text('*')
        self.emphasis = True

    def depart_emphasis(self, node):
        dprint()
        # self.add_text('*')
        self.emphasis = False

    def visit_title(self, node):
        dprint()
        self.current_paragraph = self.current_state.location.add_heading(
            level=self.sectionlevel)

    depart_title = just_print

    def visit_figure(self, node):
        # FIXME: figure text become normal paragraph instead of caption.
        dprint()

        logger.info('ATTRIBUTES:FIGURE:{}'.format(node.attributes))
        self.in_figure = True
        curloc = self.current_state.location
        logger.info('CURRENT_LOC:FIGURE:{}'.format(repr(curloc)))

        if 'align' in node.attributes:
            if node.attributes['align'] == 'center':
                self.center = True

    def depart_figure(self, node):
        dprint()
        self.in_figure = False
        if 'align' in node.attributes and 'center' == node.attributes['align']:
            self.center = False

    def visit_caption(self, node):
        dprint()

        logger.info('ATTRIBUTES:FIGURE:{}'.format(repr(node.attlist())))
        curloc = self.current_state.location

        contents = None
        if self.in_figure:
            contents = self.current_state.next_figure_num
            self.current_state.next_figure_num += 1

        self.current_paragraph = self.add_paragraph(curloc, 'Figure ', style='Caption')
        self.add_seq_field('Figure', contents)
        self.current_paragraph.add_run(': ')

    def depart_caption(self, node):
        self.current_paragraph = self.add_paragraph(self.current_state.location)

    def visit_tabular_col_spec(self, node):
        dprint()
        # TODO: properly implement this!!
        spec = node['spec']
        widths = [float(l.split('cm')[0]) for l in spec.split("{")[1:]]
        self.current_state.column_widths = widths
        raise nodes.SkipNode

    def visit_colspec(self, node):
        dprint()
        # The difficulty here is getting the right column width.
        # This can be specified with a tabular_col_spec, see above.
        #
        # Otherwise it is derived from the number of columns, which is
        # defined in visit_tgroup (a bit hackish).
        # The _block_width is the full width of the document, and this
        # is divided by the number of columns.
        #
        # It would perhaps also be possible to use node['colwidth'] in some way.
        # node['colwidth'] contains an integer like 22, the width of the column in ascii
        if self.current_state.column_widths:
            width = self.current_state.column_widths[0]
            self.current_state.column_widths = self.current_state.column_widths[1:]
            col = self.current_state.table.add_column(Cm(width))
        else:
            # noinspection PyProtectedMember
            col = self.current_state.table.add_column(
                self.docx_container._block_width // self.current_state.ncolumns)

        raise nodes.SkipNode

    depart_colspec = just_print

    def visit_tgroup(self, node):
        dprint()
        colspecs = [c for c in node.children if isinstance(c, nodes.colspec)]
        self.current_state.ncolumns = len(colspecs)

    def depart_tgroup(self, node):
        dprint()
        self.current_state.ncolumns = 1
        pass

    def visit_row(self, node):
        dprint()
        self.current_state.row = self.current_state.table.add_row()
        self.current_state.cell_counter = 0

    depart_row = just_print

    def visit_entry(self, node):
        dprint()
        if 'morerows' in node:
            raise NotImplementedError('Row spanning cells are not implemented.')
        if 'morecols' in node:
            # Hack to make column spanning possible. TODO FIX
            self.current_state.more_cols = node['morecols']

        cell = self.current_state.row.cells[self.current_state.cell_counter]
        # A new paragraph will be added by Sphinx, so remove the automated one
        # This turns out to be not possible, so instead the existing one is
        # reused in visit_paragraph.
        # cell.paragraphs.pop()
        if self.current_state.more_cols:
            # Perhaps this commented line works no too.
            # cell = cell.merge(self.row.cells[self.cell_counter + self.more_cols])
            for i in range(self.current_state.more_cols):
                cell = cell.merge(self.current_state.row.cells[
                                      self.current_state.cell_counter + i + 1])

        self.new_state(location=cell)
        # For some annoying reason, a new paragraph is automatically added
        # to each table cell. This is frustrating when you want, e.g. to
        # add a list item instead of a normal paragraph.
        self.current_paragraph = cell.paragraphs[0]

    def depart_entry(self, node):
        dprint()
        self.end_state()
        self.current_state.cell_counter = self.current_state.cell_counter + self.current_state.more_cols + 1
        self.current_state.more_cols = 0

    def visit_table(self, node):
        dprint()

        style = self.current_state.table_style
        try:
            # Check whether the style is part of the document.
            self.docx_container.styles.get_style_id(style, WD_STYLE_TYPE.TABLE)
        except KeyError as exc:
            msg = 'looks like style "{}" is missing\n{}\n using no style'.format(
                style, repr(exc))
            logger.warning(msg)
            style = None

        # Columns are added when a colspec is visited.

        # It is only possible to use a style in add_table when adding a
        # table to the root document. That is, not for a table in a table.
        if len(self.old_states):
            self.current_state.table = self.current_state.location.add_table(rows=0,
                                                                             cols=0)
        else:
            self.current_state.table = self.current_state.location.add_table(
                rows=0, cols=0, style=style)

    def depart_table(self, node):
        dprint()

        self.current_state.table = None
        self.current_state.table_style = self.table_style_default

        # Add an empty paragraph to prevent tables from being concatenated.
        # TODO: Figure out some better solution.
        self.add_paragraph(self.current_state.location, "")

    def visit_Text(self, node):
        dprint()
        text = node.astext()
        if not self.in_literal_block:
            # assert '\n\n' not in text, 'Found \n\n'
            # Replace double enter with single enter, and single enter with space.
            string_magic = 'TWOENTERSMAGICSTRING'
            text = text.replace('\n\n', string_magic)
            text = text.replace('\n', ' ')
            text = text.replace(string_magic, '\n')
        self.add_text(text)

    depart_Text = just_print

    def visit_image(self, node):
        dprint()
        logger.info('ATTRIBUTES:FIGURE:{}'.format(repr(node.attributes)))
        uri = node.attributes['uri']
        file_path = os.path.join(self.builder.env.srcdir, uri)
        dest = self.docx_container
        if self.in_figure:
            self.current_paragraph = self.add_paragraph(self.current_state.location)
            dest = self.current_paragraph.add_run()

        width = None
        height = None

        if 'height' in node.attributes:
            height = Inches(float(node.attributes['height'][0:-2]))

        if 'width' in node.attributes:
            width = Inches(float(node.attributes['width'][0:-2]))

        logger.info("width: {}, height: {}".format(width, height))
        dest.add_picture(file_path, height=height, width=width)
        # .. todo:: 'width' keyword is not supported

    depart_image = just_print

    def visit_bullet_list(self, node):
        dprint()
        # TODO: Apparently it is necessary to take into account whether
        # the list is numbered or not, like the original code did.
        # But that code did not properly account for the level.
        # So merge these two attempts.
        # self.list_style.append('ListBullet')
        self.list_level += 1

    def depart_bullet_list(self, node):
        dprint()
        # TODO: self.list_style.pop()
        self.list_level -= 1

    def visit_enumerated_list(self, node):
        dprint()
        # TODO: self.list_style.append('ListNumber')
        self.list_level += 1

    def depart_enumerated_list(self, node):
        dprint()
        # TODO: self.list_style.pop()
        self.list_level -= 1

    def visit_list_item(self, node):
        dprint()
        # A new paragraph is created here, but the next visit is to
        # paragraph, so that would add another paragraph. That is
        # prevented if current_paragraph is an empty List paragraph.
        style = 'List Bullet' if self.list_level < 2 else 'List Bullet {}'.format(
            self.list_level)
        try:
            # Check whether the style is part of the document.
            self.docx_container.styles.get_style_id(style, WD_STYLE_TYPE.PARAGRAPH)
        except KeyError as exc:
            msg = 'looks like style "{}" is missing\n{}\n using no style'.format(
                style, repr(exc))
            logger.warning(msg)
            style = None

        curloc = self.current_state.location
        if isinstance(curloc, _Cell):
            if len(curloc.paragraphs) == 1:
                if not curloc.paragraphs[0].text:
                    # An empty paragraph is created when a Cell is created.
                    # Reuse this paragraph.
                    self.current_paragraph = curloc.paragraphs[0]
                    self.current_paragraph.style = style
                else:
                    self.current_paragraph = self.add_paragraph(curloc, style=style)
            else:
                self.current_paragraph = self.add_paragraph(curloc, style=style)
        else:
            self.current_paragraph = self.add_paragraph(curloc, style=style)

    depart_list_item = just_print

    def visit_paragraph(self, node):
        dprint()

        curloc = self.current_state.location

        if 'List' in self.current_paragraph.style.name and not self.current_paragraph.text:
            # This is the first paragraph in a list item, so do not create another one.
            pass
        elif isinstance(curloc, _Cell):
            if len(curloc.paragraphs) == 1:
                if not curloc.paragraphs[0].text:
                    # An empty paragraph is created when a Cell is created.
                    # Reuse this paragraph.
                    self.current_paragraph = curloc.paragraphs[0]
                else:
                    self.current_paragraph = self.add_paragraph(curloc)
            else:
                self.current_paragraph = self.add_paragraph(curloc)
            # HACK because the style is messed up, TODO FIX
            self.current_paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            self.current_paragraph.paragraph_format.left_indent = 0
        else:
            self.current_paragraph = self.add_paragraph(curloc)

    depart_paragraph = just_print

    def visit_literal_block(self, node):
        dprint()
        # TODO: Check whether literal blocks work in tables and lists.
        self.in_literal_block = True

        logger.info('ATTRIBUTES::{}'.format(repr(node.attlist())))

        # Unlike with Lists, there will not be a visit to paragraph in a
        # literal block, so we *must* create the paragraph here.
        style = 'Preformatted Text'
        try:
            # Check whether the style is part of the document.
            self.docx_container.styles.get_style_id(style, WD_STYLE_TYPE.PARAGRAPH)
        except KeyError as exc:
            msg = 'looks like style "{}" is missing\n{}\n using no style'.format(
                style, repr(exc))
            logger.warning(msg)
            style = None

        self.current_paragraph = self.add_paragraph(self.current_state.location, style=style)
        self.current_paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    def depart_literal_block(self, node):
        dprint()
        self.in_literal_block = False

    ######## UNIMPLEMENTED NODES
    depart_start_of_file = just_print

    visit_document = just_print
    depart_document = just_print

    visit_highlightlang = print_and_skip

    visit_topic = print_and_skip
    depart_topic = print_and_skip

    visit_sidebar = visit_topic
    depart_sidebar = depart_topic

    visit_rubric = print_and_skip
    # self.add_text('-[ ')

    depart_rubric = print_and_skip
    # self.add_text(' ]-')

    visit_compound = just_print
    depart_compound = just_print

    visit_glossary = just_print
    depart_glossary = just_print

    visit_subtitle = just_print
    depart_subtitle = just_print

    visit_attribution = print_and_skip
    # self.add_text('-- ')
    depart_attribution = just_print

    visit_desc = just_print
    depart_desc = print_and_skip

    visit_desc_signature = print_and_skip
    depart_desc_signature = print_and_skip

    visit_desc_name = just_print
    depart_desc_name = just_print

    visit_desc_addname = just_print
    depart_desc_addname = just_print

    visit_desc_type = just_print
    depart_desc_type = just_print

    visit_desc_returns = print_and_skip
    # self.add_text(' -> ')
    depart_desc_returns = just_print

    visit_desc_parameterlist = print_and_skip
    # self.add_text('(')
    # self.first_param = 1

    depart_desc_parameterlist = print_and_skip
    # self.add_text(')')

    visit_desc_parameter = print_and_skip
    # if not self.first_param:
    #     self.add_text(', ')
    # else:
    #     self.first_param = 0
    # self.add_text(node.astext())
    # raise nodes.SkipNode

    depart_desc_parameter = print

    visit_desc_optional = print_and_skip
    # self.add_text('[')

    depart_desc_optional = print_and_skip
    # self.add_text(']')

    visit_desc_annotation = just_print
    depart_desc_annotation = just_print

    visit_refcount = just_print
    depart_refcount = just_print

    visit_desc_content = print_and_skip
    depart_desc_content = print_and_skip

    def visit_productionlist(self, node):
        dprint()
        raise nodes.SkipNode
        # names = []
        # for production in node:
        #     names.append(production['tokenname'])
        # maxlen = max(len(name) for name in names)
        # for production in node:
        #     if production['tokenname']:
        #         self.add_text(production['tokenname'].ljust(maxlen) + ' ::=')
        #         lastname = production['tokenname']
        #     else:
        #         self.add_text('%s    ' % (' '*len(lastname)))
        #     self.add_text(production.astext() + '\n')
        # raise nodes.SkipNode

    visit_seealso = just_print
    depart_seealso = just_print

    def visit_footnote(self, node):
        dprint()
        raise nodes.SkipNode
        # self._footnote = node.children[0].astext().strip()

    depart_footnote = print_and_skip

    def visit_citation(self, node):
        dprint()
        raise nodes.SkipNode
        # if len(node) and isinstance(node[0], nodes.label):
        #     self._citlabel = node[0].astext()
        # else:
        #     self._citlabel = ''

    depart_citation = print_and_skip

    def visit_label(self, node):
        dprint()
        raise nodes.SkipNode

    # XXX: option list could use some better styling

    visit_option_list = just_print

    def depart_option_list(self, node):
        dprint()
        pass

    visit_option_list_item = print_and_skip
    depart_option_list_item = print_and_skip

    def visit_option_group(self, node):
        dprint()
        raise nodes.SkipNode
        # self._firstoption = True

    def depart_option_group(self, node):
        dprint()
        raise nodes.SkipNode
        # self.add_text('     ')

    def visit_option(self, node):
        dprint()
        raise nodes.SkipNode
        # if self._firstoption:
        #     self._firstoption = False
        # else:
        #     self.add_text(', ')

    depart_option = just_print

    def visit_option_string(self, node):
        dprint()
        pass

    depart_option_string = just_print

    def visit_option_argument(self, node):
        dprint()
        raise nodes.SkipNode
        # self.add_text(node['delimiter'])

    depart_option_argument = just_print

    visit_description = just_print
    depart_description = just_print

    visit_thead = just_print
    depart_thead = just_print

    visit_tbody = just_print
    depart_tbody = just_print

    def visit_acks(self, node):
        dprint()
        raise nodes.SkipNode
        # self.add_text(', '.join(n.astext() for n in node.children[0].children)
        #               + '.')

    def visit_transition(self, node):
        dprint()
        raise nodes.SkipNode
        # self.add_text('=' * 70)

    def visit_definition_list(self, node):
        dprint()
        raise nodes.SkipNode
        # self.list_style.append(-2)

    def depart_definition_list(self, node):
        dprint()
        raise nodes.SkipNode
        # self.list_style.pop()

    visit_definition_list_item = print_and_skip

    def depart_definition_list_item(self, node):
        dprint()
        pass

    visit_term = print_and_skip

    depart_term = print_and_skip

    def visit_classifier(self, node):
        dprint()
        raise nodes.SkipNode
        # self.add_text(' : ')

    depart_classifier = print_and_skip

    visit_definition = print_and_skip

    depart_definition = print_and_skip

    visit_field_list = just_print

    def depart_field_list(self, node):
        dprint()
        pass

    visit_field = just_print

    def depart_field(self, node):
        dprint()
        pass

    visit_field_name = print_and_skip

    def depart_field_name(self, node):
        dprint()
        raise nodes.SkipNode
        # self.add_text(':')

    visit_field_body = print_and_skip

    depart_field_body = print_and_skip

    visit_centered = just_print

    def depart_centered(self, node):
        dprint()
        pass

    visit_hlist = just_print

    def depart_hlist(self, node):
        dprint()
        pass

    visit_hlistcol = just_print

    def depart_hlistcol(self, node):
        dprint()
        pass

    visit_admonition = print_and_skip

    depart_admonition = print_and_skip

    _visit_admonition = print_and_skip

    visit_attention = _visit_admonition
    depart_attention = _make_depart_admonition('attention')
    visit_caution = _visit_admonition
    depart_caution = _make_depart_admonition('caution')
    visit_danger = _visit_admonition
    depart_danger = _make_depart_admonition('danger')
    visit_error = _visit_admonition
    depart_error = _make_depart_admonition('error')
    visit_hint = _visit_admonition
    depart_hint = _make_depart_admonition('hint')
    visit_important = _visit_admonition
    depart_important = _make_depart_admonition('important')
    visit_note = _visit_admonition
    depart_note = _make_depart_admonition('note')
    visit_tip = _visit_admonition
    depart_tip = _make_depart_admonition('tip')
    visit_warning = _visit_admonition
    depart_warning = _make_depart_admonition('warning')

    def visit_versionmodified(self, node):
        dprint()
        raise nodes.SkipNode
        # from sphinx.locale import admonitionlabels, versionlabels, _
        # if node.children:
        #     self.add_text(
        #             versionlabels[node['type']] % node['version'] + ': ')
        # else:
        #     self.add_text(
        #             versionlabels[node['type']] % node['version'] + '.')

    depart_versionmodified = print_and_skip

    visit_doctest_block = print_and_skip

    depart_doctest_block = print_and_skip

    visit_line_block = print_and_skip

    depart_line_block = print_and_skip

    visit_line = just_print

    def depart_line(self, node):
        dprint()
        pass

    visit_block_quote = just_print

    def depart_block_quote(self, node):
        dprint()

    visit_compact_paragraph = just_print

    depart_compact_paragraph = just_print

    visit_target = print_and_skip

    visit_index = print_and_skip

    visit_substitution_definition = print_and_skip

    visit_pending_xref = just_print
    depart_pending_xref = just_print

    visit_reference = just_print
    depart_reference = just_print

    visit_download_reference = just_print
    depart_download_reference = just_print

    def visit_literal_emphasis(self, node):
        dprint()
        # self.add_text('*')

    def depart_literal_emphasis(self, node):
        dprint()
        # self.add_text('*')

    def visit_abbreviation(self, node):
        dprint()
        # self.add_text('')

    def depart_abbreviation(self, node):
        dprint()
        # if node.hasattr('explanation'):
        #     self.add_text(' (%s)' % node['explanation'])

    def visit_title_reference(self, node):
        dprint()
        # self.add_text('*')

    def depart_title_reference(self, node):
        dprint()
        # self.add_text('*')

    def visit_literal(self, node):
        dprint()
        # self.add_text('``')

    def depart_literal(self, node):
        dprint()
        # self.add_text('``')

    def visit_subscript(self, node):
        dprint()
        raise nodes.SkipNode
        # self.add_text('_')

    depart_subscript = just_print

    def visit_superscript(self, node):
        dprint()
        raise nodes.SkipNode
        # self.add_text('^')

    depart_superscript = just_print

    def visit_footnote_reference(self, node):
        dprint()
        raise nodes.SkipNode
        # self.add_text('[%s]' % node.astext())

    def visit_citation_reference(self, node):
        dprint()
        raise nodes.SkipNode
        # self.add_text('[%s]' % node.astext())

    def visit_generated(self, node):
        dprint()
        pass

    depart_generated = just_print

    def visit_inline(self, node):
        dprint()
        pass

    depart_inline = just_print

    def visit_problematic(self, node):
        dprint()
        raise nodes.SkipNode
        # self.add_text('>>')

    def depart_problematic(self, node):
        dprint()
        raise nodes.SkipNode
        # self.add_text('<<')

    def visit_system_message(self, node):
        dprint()
        raise nodes.SkipNode
        # self.add_text('<SYSTEM MESSAGE: %s>' % node.astext())

    def visit_meta(self, node):
        dprint()
        raise nodes.SkipNode
        # only valid for HTML

    def visit_raw(self, node):
        dprint()
        raise nodes.SkipNode
        # if 'text' in node.get('format', '').split():
        #     self.body.append(node.astext())

    visit_container = just_print
    depart_container = just_print

    def unknown_visit(self, node):
        dprint()
        raise nodes.SkipNode
        # raise NotImplementedError('Unknown node: ' + node.__class__.__name__)

    def unknown_departure(self, node):
        dprint()
        raise nodes.SkipNode
        # raise NotImplementedError('Unknown node: ' + node.__class__.__name__)
